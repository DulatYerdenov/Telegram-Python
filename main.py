import asyncio
import logging
import sqlite3
import openai
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import subprocess
import os
from aiogram import F


from aiogram import Bot, Dispatcher, types
from aiogram.filters import Command
from aiogram.enums import ContentType
from aiogram.types import BotCommand, BotCommandScopeDefault, ReplyKeyboardMarkup, KeyboardButton
from aiogram.utils.markdown import hbold

from dotenv import load_dotenv

load_dotenv()

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Токен вашего бота
TOKEN_BOT = os.getenv('TOKEN_BOT')

# ID администратора
ADMIN_ID = os.getenv('ADMIN_ID')

# Инициализация бота и диспетчера с увеличенным временем ожидания
bot = Bot(token=TOKEN_BOT, timeout=30)  # Увеличиваем тайм-аут до 30 секунд
dp = Dispatcher()

# Подключение к базе данных SQLite
conn = sqlite3.connect('my_database.db')
cursor = conn.cursor()

# Создание таблицы users, если она не существует
cursor.execute('''
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY,
        user_id INTEGER UNIQUE,
        first_name TEXT,
        last_name TEXT,
        username TEXT
    )
''')
conn.commit()

# Настройка OpenAI
client = openai.OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

# Константа для оценки времени обработки одной задачи (в секундах)
TIME_PER_TASK = 5  
MINIMUM_TIME = 5

# Команды бота
async def set_commands(bot: Bot):
    commands = [
        BotCommand(command="start", description="Начало работы"),
        BotCommand(command="help", description="Помощь"),
        BotCommand(command="lesson", description="Получить урок"),
        BotCommand(command="quiz", description="Пройти тест"),
        BotCommand(command="project", description="Идеи для проектов"),
        BotCommand(command="export", description="Выгрузить данные"),
        BotCommand(command="ask", description="Задать вопрос по Python"),
    ]
    await bot.set_my_commands(commands, scope=BotCommandScopeDefault())

# Создаем клавиатуру
buttons = [
    KeyboardButton(text='/start'),
    KeyboardButton(text='/help'),
    KeyboardButton(text='/lesson'),
    KeyboardButton(text='/quiz'),
    KeyboardButton(text='/project'),
    KeyboardButton(text='/export'),
    KeyboardButton(text='/ask'),
]
kb = ReplyKeyboardMarkup(keyboard=[buttons], resize_keyboard=True)

# Обработчик команды /start
@dp.message(Command("start"))
async def cmd_start(message: types.Message):
    user_id = message.from_user.id
    cursor.execute("SELECT * FROM users WHERE user_id=?", (user_id,))
    user = cursor.fetchone()

    if user is None:
        cursor.execute("INSERT INTO users (user_id, first_name, last_name, username) VALUES (?, ?, ?, ?)",
                       (user_id, message.from_user.first_name, message.from_user.last_name, message.from_user.username))
        conn.commit()
        await message.answer(f"Приветствую, {hbold(message.from_user.first_name)}! \n"
                             "Вы успешно зарегистрированы!", reply_markup=kb)
    else:
        await message.answer(f"Приветствую, {hbold(message.from_user.first_name)}! \n"
                             "Рад видеть вас снова!", reply_markup=kb)

    await message.answer(
        "Я — бот для изучения Python. Вот что я умею:\n"
        "/start - Начать работу\n"
        "/help - Получить помощь\n"
        "/lesson - Получить урок\n"
        "/quiz - Пройти тест\n"
        "/project - Идеи для проектов\n"
        "/ask - Задать вопрос по Python\n"
        "Или отправьте файл .doc/.docx с задачами, чтобы я решил их!"
    )

# Обработчик команды /help
@dp.message(Command("help"))
async def cmd_help(message: types.Message):
    await message.answer(
        "Список доступных команд:\n"
        "/start - Начать работу\n"
        "/help - Получить помощь\n"
        "/lesson - Получить урок\n"
        "/quiz - Пройти тест\n"
        "/project - Идеи для проектов\n"
        "/export - Выгрузить данные (только для администратора)\n"
        "/ask - Задать вопрос по Python\n"
        "Или отправьте файл .doc/.docx с задачами, чтобы я решил их!"
    )

# Обработчик команды /lesson
@dp.message(Command("lesson"))
async def cmd_lesson(message: types.Message):
    await message.answer("Вот ваш урок: [https://metanit.com/python/tutorial/]")

# Обработчик команды /quiz
@dp.message(Command("quiz"))
async def cmd_quiz(message: types.Message):
    await message.answer("Вот ваш тест: [https://docs.google.com/document/d/1Y1YU1lFDPkYeXshVAESStlRW5AvXb_VpMDRGWSXDR0o/edit?usp=sharing]")

# Обработчик команды /project
@dp.message(Command("project"))
async def cmd_project(message: types.Message):
    projects = [
        "1. Калькулятор",
        "2. To-Do приложение",
        "3. Чат-бот",
        "4. Парсер веб-страниц",
        "5. Игра 'Угадай число'",
        "6. Конвертер валют",
        "7. Погодное приложение",
        "8. Генератор паролей"
    ]
    await message.answer("Идеи для проектов:\n" + "\n".join(projects))

# Обработчик команды /export (только для администратора)
@dp.message(Command("export"))
async def cmd_export(message: types.Message):
    if str(message.from_user.id) == ADMIN_ID:  # Проверяем ID пользователя
        try:
            cursor.execute("SELECT * FROM users")
            users = cursor.fetchall()

            if users:
                with open("users_data.txt", "w") as file:
                    for user in users:
                        file.write(f"ID: {user[0]}, User ID: {user[1]}, Имя: {user[2]}, Фамилия: {user[3]}, Username: {user[4]}\n")
                await message.answer_document(types.FSInputFile("users_data.txt"))
            else:
                await message.answer("В базе данных нет пользователей.")

        except Exception as e:
            logger.error(f"Ошибка при выгрузке данных: {e}")
            await message.answer("Произошла ошибка при выгрузке данных.")
    else:
        await message.answer("У вас нет прав администратора для выполнения этой команды.")

# Обработчик команды /ask (вопросы к OpenAI)
@dp.message(Command("ask"))
async def cmd_ask(message: types.Message):
    await message.answer("Напишите ваш вопрос по Python, и я постараюсь на него ответить.")

# Обработчик текстовых сообщений (для вопросов к OpenAI)
@dp.message(F.text)
async def handle_message(message: types.Message):
    if message.text.startswith("/"):
        return  # Игнорируем команды

    # Отправляем вопрос в OpenAI
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Ты помощник, который отвечает на вопросы по Python."},
                {"role": "user", "content": f"Ответь на вопрос по Python: {message.text}"}
            ],
            max_tokens=500,
            temperature=0.7,
        )
        answer = response.choices[0].message.content.strip()
        await message.answer(answer)
    except Exception as e:
        logger.error(f"Ошибка при обращении к OpenAI: {e}")
        await message.answer("Произошла ошибка при обработке вашего вопроса.")

# Обработчик для документов (.doc и .docx)
@dp.message(F.content_type == ContentType.DOCUMENT)
async def handle_document(message: types.Message):
    # Проверяем расширение файла
    file_name = message.document.file_name
    if not (file_name.endswith('.doc') or file_name.endswith('.docx')):
        await message.answer("Пожалуйста, отправьте файл в формате .doc или .docx.")
        return

    try:
        # Скачиваем файл
        file_info = await bot.get_file(message.document.file_id)
        file = await bot.download_file(file_info.file_path)

        # Сохраняем файл временно
        input_file = f"input_{file_name}"
        with open(input_file, "wb") as f:
            f.write(file.read())

        # Если это .doc, конвертируем в .docx
        if file_name.endswith('.doc'):
            output_file = "converted.docx"
            try:
                subprocess.run([
                    "libreoffice", "--headless", "--convert-to", "docx",
                    input_file, "--outdir", "."
                ], check=True, timeout=30)
            except subprocess.SubprocessError as e:
                logger.error(f"Ошибка конвертации .doc в .docx: {e}")
                await message.answer("Ошибка при конвертации файла .doc в .docx.")
                return
            if not os.path.exists(output_file):
                await message.answer("Не удалось конвертировать файл .doc в .docx.")
                return
            docx_file = output_file
        else:
            docx_file = input_file

        # Читаем содержимое файла .docx
        doc = Document(docx_file)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():  # Игнорируем пустые абзацы
                full_text.append(para.text)
        if not full_text:
            await message.answer("Документ пустой. Пожалуйста, отправьте документ с примерами по Python.")
            return

        # Разделяем текст на задачи (предполагаем, что каждая задача начинается с "Вопрос")
        tasks = []
        current_task = []
        for line in full_text:
            if line.startswith("Вопрос"):
                if current_task:  # Сохраняем предыдущую задачу
                    tasks.append("\n".join(current_task))
                current_task = [line]  # Начинаем новую задачу
            else:
                current_task.append(line)
        if current_task:  # Не забываем последнюю задачу
            tasks.append("\n".join(current_task))

        if not tasks:
            await message.answer("Не найдено ни одной задачи. Убедитесь, что задачи начинаются с 'Вопрос'.")
            return

        # Оцениваем время обработки
        num_tasks = len(tasks)
        estimated_time = max(MINIMUM_TIME, num_tasks * TIME_PER_TASK)  # Минимальное время или время на задачи
        await message.answer(f"Ваш документ обрабатывается. Примерное время готовности: {estimated_time} секунд.")

        # Обрабатываем каждую задачу через OpenAI
        solutions = []
        for task in tasks:
            try:
                response = client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": "Ты помощник, который решает задачи по Python. Дай подробное решение задачи."},
                        {"role": "user", "content": f"Реши задачу по Python:\n{task}"}
                    ],
                    max_tokens=500,
                    temperature=0.7,
                )
                answer = response.choices[0].message.content.strip()
                solutions.append((task, answer))
            except Exception as e:
                logger.error(f"Ошибка при обработке задачи: {e}")
                solutions.append((task, "Ошибка при решении задачи."))

        # Создаем новый документ .docx с задачами и ответами
        output_doc = Document()
        output_doc.add_heading("Задачи и решения", level=1)

        for task, solution in solutions:
            # Добавляем заголовок задачи
            output_doc.add_heading(task.split('\n')[0], level=2)

            # Добавляем задачу
            output_doc.add_paragraph(task + "\n")

            # Добавляем решение с выделением красным цветом
            solution_para = output_doc.add_paragraph()
            solution_run = solution_para.add_run(solution)
            solution_run.font.color.rgb = RGBColor(255, 0, 0)  # Красный цвет
            solution_run.font.size = Pt(12)

            # Проверяем, существует ли w:pPr, и если нет, создаем его
            paragraph = solution_run._element.getparent()
            pPr = paragraph.xpath('.//w:pPr')
            if not pPr:  # Если w:pPr не существует, добавляем его
                pPr_element = OxmlElement('w:pPr')
                paragraph.insert(0, pPr_element)
                pPr = paragraph.xpath('.//w:pPr')

            # Создаем элемент w:pBdr для рамки
            pBdr = OxmlElement('w:pBdr')
            for border in ['top', 'bottom', 'left', 'right']:
                border_element = OxmlElement(f'w:{border}')
                border_element.set(qn('w:val'), 'single')
                border_element.set(qn('w:sz'), '4')
                border_element.set(qn('w:space'), '1')
                border_element.set(qn('w:color'), 'FF0000')
                pBdr.append(border_element)
            pPr[0].append(pBdr)

            # Добавляем отступы (две пустые строки)
            output_doc.add_paragraph("")
            output_doc.add_paragraph("")

        # Сохраняем документ в файл
        output_file = "answers.docx"
        output_doc.save(output_file)

        # Отправляем документ пользователю
        await message.answer_document(types.FSInputFile(output_file))

    except Exception as e:
        logger.error(f"Ошибка при обработке документа: {e}")
        await message.answer("Произошла ошибка при обработке документа.")
    finally:
        # Удаляем временные файлы, если они были созданы
        for temp_file in [input_file, "converted.docx", "answers.docx"]:
            if os.path.exists(temp_file):
                os.remove(temp_file)

# Запуск бота с увеличенным количеством попыток подключения
async def main():
    await set_commands(bot)
    max_retries = 5  # Максимальное количество попыток подключения
    for attempt in range(max_retries):
        try:
            await dp.start_polling(bot)
            break
        except Exception as e:
            logger.error(f"Ошибка подключения (попытка {attempt + 1}/{max_retries}): {e}")
            if attempt < max_retries - 1:
                await asyncio.sleep(2 ** attempt)  # Экспоненциальная задержка перед повторной попыткой
            else:
                logger.error("Не удалось подключиться после всех попыток. Завершение работы.")
                raise

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        logger.info("Бот остановлен.")
    finally:
        conn.close()  # Закрываем соединение с базой данных при завершении работы